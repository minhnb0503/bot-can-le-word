import os
import re
import telebot
import logging
from threading import Thread
from flask import Flask
from docx import Document
from docx.shared import Mm, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# --- 1. HỆ THỐNG GHI LOGS (BẮT LỖI) ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)

# --- 2. CẤU HÌNH TOKEN TELEGRAM CỦA BẠN ---
API_TOKEN = '8772669126:AAFRYcViBI3dG_2MfDkA6VZtGIi-nSm0eBM'
bot = telebot.TeleBot(API_TOKEN)

# --- 3. TẠO WEB SERVER PHỤ (Giúp giữ bot luôn chạy online trên Render) ---
app = Flask('')

@app.route('/')
def home():
    return "Bot định dạng Word đang chạy trực tuyến 24/7!"

def run_web_server():
    port = int(os.environ.get("PORT", 8080))
    app.run(host='0.0.0.0', port=port)

# --- 4. HÀM ĐỊNH DẠNG FONT TIMES NEW ROMAN ---
def set_run_font_times_new_roman(run, size_pt=None, bold=None):
    """Ép font chuẩn Times New Roman. Nếu không truyền cỡ chữ/đậm thì giữ nguyên của file gốc"""
    run.font.name = 'Times New Roman'
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
        
    # Ép cấu trúc XML Word nhận diện chuẩn Unicode tiếng Việt
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# --- 5. HÀM XỬ LÝ ĐỊNH DẠNG TOÀN DIỆN BÁO CÁO ---
def format_docx(input_path, output_path):
    logging.info(f"Bắt đầu xử lý cấu trúc file: {input_path}")
    doc = Document(input_path)
    
    # Cấu hình khổ giấy A4 và Căn lề chuẩn: Trái 30mm, Phải 15mm, Trên 20mm, Dưới 20mm
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)

    # Biến đánh dấu: Khi nào chưa gặp Mục lục/Lời mở đầu/Chương 1 thì CHƯA kích hoạt căn lề để bảo vệ trang bìa
    start_strict_formatting = False
    
    # Từ khóa nhận diện các phần bắt buộc phải sang trang mới tinh
    major_sections = ["MỤC LỤC", "LỜI MỞ ĐẦU", "KẾT LUẬN", "TÀI LIỆU THAM KHẢO", "PHỤ LỤC", "DANH MỤC"]

    for paragraph in doc.paragraphs:
        # Xóa bỏ các ký tự xuống dòng ép buộc (\v, \n) làm dãn chữ tai hại
        for run in paragraph.runs:
            if run.text:
                run.text = run.text.replace('\v', ' ').replace('\n', ' ')
                
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Kiểm tra xem đã hết Trang bìa + Bảng phân công nhiệm vụ chưa để bật định dạng báo cáo
        text_upper = text.upper()
        if any(sec in text_upper for sec in major_sections) or text_upper.startswith("CHƯƠNG"):
            start_strict_formatting = True

        # GIAI ĐOẠN 1: NẾU ĐANG Ở TRANG BÌA HOẶC BẢNG NHIỆM VỤ -> CHỈ ĐỔI FONT, GIỮ NGUYÊN CỠ CHỮ VÀ VỊ TRÍ
        if not start_strict_formatting:
            for run in paragraph.runs:
                set_run_font_times_new_roman(run) # Giữ nguyên định dạng gốc của bìa
            continue

        # GIAI ĐOẠN 2: ĐÃ VÀO NỘI DUNG BÁO CÁO CHÍNH -> ÁP DỤNG QUY TẮC NGHIÊM NGẶT
        is_chapter = False
        # Kiểm tra xem đoạn này có phải Tiêu đề lớn (Chương hoặc Mục lục...) hay không
        if text_upper.startswith("CHƯƠNG") or any(text_upper == sec for sec in major_sections):
            is_chapter = True

        # Kiểm tra xem có phải mục con dạng số (1., 1.1, 2.1.1...) hay không
        is_section = False
        if not is_chapter and re.match(r'^(\d+(\.\d+)*\.?)\s+', text):
            is_section = True

        # ÁP ĐỊNH DẠNG CHI TIẾT THEO TỪNG THÀNH PHẦN
        if is_chapter:
            # TỰ ĐỘNG NGẮT TRANG: Ép phần này phải nhảy sang đầu trang mới
            paragraph.paragraph_format.page_break_before = True
            
            # Định dạng: Chữ in hoa, cao 16, Đậm, căn lề giữa
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.text = text.upper()
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = 1.3
            for run in paragraph.runs:
                set_run_font_times_new_roman(run, size_pt=16, bold=True)
                
        elif is_section:
            # Mục con 1.1, 1.2...: Chữ thường, cỡ 13, Đậm, căn đều Justify (Bảo hiểm dòng ngắn tránh dãn chữ)
            paragraph.paragraph_format.page_break_before = False
            if len(text) < 60:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = 1.3
            for run in paragraph.runs:
                set_run_font_times_new_roman(run, size_pt=13, bold=True)
                
        else:
            # Văn bản thường: Cỡ 13, Justify, thụt dòng đầu 1cm, giãn dòng 1.3, space before 6pt
            paragraph.paragraph_format.page_break_before = False
            if len(text) < 60:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
            paragraph.paragraph_format.first_line_indent = Cm(1)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = 1.3
            for run in paragraph.runs:
                set_run_font_times_new_roman(run, size_pt=13, bold=False)

    # Duyệt qua toàn bộ bảng biểu (Tables) có trong file để ép font Times New Roman chuẩn luôn
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_font_times_new_roman(run)

    doc.save(output_path)
    logging.info(f"Đã xuất file hoàn mỹ: {output_path}")

# --- 6. ĐIỀU HƯỚNG BOT TELEGRAM ---
@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    welcome_text = (
        "👋 Xin chào! Tôi là Bot tự động căn lề báo cáo Đồ án Viễn thông (Bản siêu nâng cấp bảo vệ Trang Bìa).\n\n"
        "📥 Hãy gửi cho tôi file Word (.docx) của bạn. Tôi sẽ tự động sửa chuẩn chỉnh lề lối từ đầu đến cuối."
    )
    bot.reply_to(message, welcome_text)

@bot.message_handler(content_types=['document'])
def handle_docs(message):
    input_filename = ""
    output_filename = ""
    try:
        file_name = message.document.file_name
        logging.info(f"Yêu cầu xử lý từ Chat ID: {message.chat.id}, file: {file_name}")
        
        if not file_name.endswith('.docx'):
            bot.reply_to(message, "❌ Vui lòng chỉ gửi file định dạng Word (.docx) thôi nhé bạn ơi!")
            return
            
        bot.reply_to(message, "🔄 Bot đang phân tích cấu trúc, giữ nguyên trang bìa và tự động nhảy trang chuẩn chỉnh. Vui lòng đợi...")
        
        file_info = bot.get_file(message.document.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        
        input_filename = f"input_{message.chat.id}_{file_name}"
        output_filename = f"Formatted_{file_name}"
        
        with open(input_filename, 'wb') as new_file:
            new_file.write(downloaded_file)
            
        format_docx(input_filename, output_filename)
        
        with open(output_filename, 'rb') as formatted_doc:
            bot.send_document(message.chat.id, formatted_doc, caption="✅ Đã căn chỉnh lề, bảo vệ trang bìa và ngắt trang tự động hoàn tất!")
            
    except Exception as e:
        logging.error(f"LỖI HỆ THỐNG: {str(e)}", exc_info=True)
        bot.reply_to(message, "❌ Có lỗi cấu trúc đặc biệt trong file. Bạn hãy thử lưu lại file Word rồi gửi lại nhé.")
        
    finally:
        if input_filename and os.path.exists(input_filename):
            os.remove(input_filename)
        if output_filename and os.path.exists(output_filename):
            os.remove(output_filename)

# --- 7. KÍCH HOẠT HỆ THỐNG ---
if __name__ == '__main__':
    logging.info("--- HỆ THỐNG BOT BẮT ĐẦU KÍCH HOẠT ---")
    server_thread = Thread(target=run_web_server)
    server_thread.start()
    bot.infinity_polling()
