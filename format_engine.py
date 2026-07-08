# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════╗
║         BOT CĂN CHỈNH BÁO CÁO TỰ ĐỘNG                    ║
║         Vietnamese University Report Formatter               ║
╚══════════════════════════════════════════════════════════════╝

Tự động căn chỉnh báo cáo đại học theo quy chuẩn:
  • Lề:  Trái 3cm | Phải 1.5cm | Trên/Dưới 2cm
  • Font: Times New Roman toàn bộ
  • Heading 1: 16pt, đậm, căn giữa, IN HOA
  • Heading 2/3: 13pt, đậm, căn đều
  • Nội dung: 13pt, căn đều, thụt 1cm, giãn dòng 1.3
  • Trang bìa: Nhận diện tự động, căn đều, ẩn số trang
  • Mục lục: Tự động tạo (TOC \\o "1-3"), dấu chấm + số trang
  • Bảng: Viền đầy đủ, tiêu đề căn giữa, rộng 100%
  • Xử lý tài liệu tham khảo: Thụt treo 1cm

Trường hợp xử lý:
  1. Trang bìa: Nhận diện qua từ khóa (BỘ/TRƯỜNG/KHOA/BÁO CÁO...),
     xóa dòng trống, giãn đều, giữ logo/ảnh, bảng danh sách SV.
  2. Mục lục: Kiểm tra trùng lặp, style TOC 1-3 không customStyle,
     tab stop lề phải 16.5cm với leader dots.
  3. Heading: Giữ auto-numbering chương, bỏ numbering tiêu đề đặc biệt.
  4. Đoạn trống: Xóa toàn bộ (trừ đoạn có hình ảnh).
  5. Ngắt trang thủ công: Xóa khỏi run, dùng pageBreakBefore trên Heading 1.
  6. Bảng: Bảng 0 (trang bìa) = không viền, giữ list numbering.
     Bảng khác = viền đầy đủ, header căn giữa đậm.
  7. Tài liệu tham khảo: Thụt treo (left=1cm, first=0).
  8. File đã chỉnh: Phát hiện MỤC LỤC/TOC sẵn có → không thêm lần 2.
"""

import os
import sys
import traceback
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH, WD_LINE_SPACING,
    WD_TAB_ALIGNMENT, WD_TAB_LEADER,
)
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ╔══════════════════════════════════════════════════════════════╗
# ║ CONFIGURATION — Thay đổi ở đây nếu cần                     ║
# ╚══════════════════════════════════════════════════════════════╝
FONT_NAME       = 'Times New Roman'
BODY_SIZE_PT    = 13        # Cỡ chữ nội dung
H1_SIZE_PT      = 16        # Cỡ chữ Heading 1
LINE_SPACING_MUL= 1.3       # Giãn dòng nội dung (multiple)
FIRST_INDENT_CM = 1.0       # Thụt đầu dòng (cm)

MARGIN_LEFT_MM  = 30
MARGIN_RIGHT_MM = 15
MARGIN_TOP_MM   = 20
MARGIN_BOT_MM   = 20

# Tiêu đề đặc biệt → Heading 1, không đánh số chương
SPECIAL_TITLES = {
    # Mở đầu
    'LỜI MỞ ĐẦU', 'LỜI NÓI ĐẦU', 'MỞ ĐẦU', 'PHẦN MỞ ĐẦU',
    # Kết luận
    'KẾT LUẬN', 'LỜI KẾT', 'PHẦN KẾT LUẬN',
    # Tham khảo
    'TÀI LIỆU THAM KHẢO', 'DANH MỤC TÀI LIỆU THAM KHẢO',
    # Phụ lục & danh mục
    'PHỤ LỤC', 'LỜI CẢM ƠN', 'LỜI CAM ĐOAN',
    'DANH MỤC BẢNG', 'DANH MỤC HÌNH', 'DANH MỤC BIỂU ĐỒ',
    'DANH MỤC CHỮ VIẾT TẮT', 'DANH MỤC TỪ VIẾT TẮT',
    'DANH MỤC CÁC KÝ HIỆU', 'DANH MỤC HÌNH ẢNH',
    'DANH MỤC BẢNG BIỂU', 'DANH MỤC SƠ ĐỒ',
    'TÓM TẮT', 'ABSTRACT',
}

# Từ khóa đánh dấu kết thúc trang bìa (dùng để tìm ranh giới)
BODY_MARKERS = SPECIAL_TITLES | {'MỤC LỤC'}

# Từ khóa phân loại phần tử trang bìa
HEADER_KW  = ['BỘ', 'TRƯỜNG', 'ĐẠI HỌC', 'KHOA', 'VIỆN', 'HỌC VIỆN', 'CƠ SỞ']
TITLE_KW   = ['BÁO CÁO', 'ĐỒ ÁN', 'LUẬN VĂN', 'KHÓA LUẬN', 'TIỂU LUẬN',
              'LUẬN ÁN', 'ĐỀ TÀI', 'CHUYÊN ĐỀ', 'BÀI TẬP LỚN']

# Từ khóa xác định vùng tài liệu tham khảo (thụt treo)
REF_TITLES = {'TÀI LIỆU THAM KHẢO', 'DANH MỤC TÀI LIỆU THAM KHẢO'}

# Giới hạn trang bìa
MIN_COVER_PARAS = 3
MAX_COVER_PARAS = 35

# ╔══════════════════════════════════════════════════════════════╗
# ║ HELPER FUNCTIONS                                            ║
# ╚══════════════════════════════════════════════════════════════╝

def set_fonts_xml(obj, font_name=FONT_NAME):
    """Đặt font ở tầng XML (rPr/rFonts) cho style hoặc run."""
    rPr = obj._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        obj._element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
        rFonts.set(qn(attr), font_name)


def set_run_fonts(run, font_name=FONT_NAME):
    """Đặt font cho 1 run cả ở API lẫn XML."""
    run.font.name = font_name
    set_fonts_xml(run, font_name)


def disable_numbering(para):
    """Tắt auto-numbering cho 1 đoạn cụ thể (numId=0)."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        para._element.insert(0, pPr)
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)
    pPr.append(parse_xml(f'<w:numPr {nsdecls("w")}><w:numId w:val="0"/></w:numPr>'))


def has_visual_elements(para):
    """Kiểm tra đoạn có chứa hình ảnh / drawing / OLE object không."""
    el = para._element
    return (
        len(el.findall('.//' + qn('w:drawing'))) > 0 or
        len(el.findall('.//' + qn('w:pict')))    > 0 or
        len(el.findall('.//' + qn('w:object')))  > 0
    )


def max_run_font_size(para):
    """Tìm cỡ font lớn nhất trong các run của đoạn (pt). Trả 0 nếu không xác định."""
    mx = 0
    for r in para.runs:
        if r.font.size:
            mx = max(mx, r.font.size.pt)
    return mx

# ╔══════════════════════════════════════════════════════════════╗
# ║ COVER PAGE DETECTION & CLASSIFICATION                       ║
# ╚══════════════════════════════════════════════════════════════╝

def find_cover_end(doc):
    """
    Tìm chỉ số đoạn văn đầu tiên của phần nội dung (kết thúc trang bìa).
    
    Thuật toán:
    1. Quét từ đoạn MIN_COVER_PARAS trở đi
    2. Dừng khi gặp: Heading style HOẶC tiêu đề đặc biệt (LỜI MỞ ĐẦU...)
    3. Giới hạn tối đa MAX_COVER_PARAS
    4. Nếu không tìm thấy → trả về 0 (không có trang bìa)
    """
    for i, para in enumerate(doc.paragraphs):
        if i < MIN_COVER_PARAS:
            continue
        if i >= MAX_COVER_PARAS:
            return MAX_COVER_PARAS
        style = para.style.name if para.style else ''
        text  = para.text.strip()
        # Gặp heading style → kết thúc trang bìa
        if 'Heading' in style:
            return i
        # Gặp tiêu đề đặc biệt (LỜI MỞ ĐẦU, MỤC LỤC...)
        if text in BODY_MARKERS:
            return i
    return 0  # Không xác định được → không có trang bìa


def classify_cover_element(para):
    """
    Phân loại phần tử trang bìa để đặt spacing phù hợp.
    
    Trả về 1 trong:
      'header', 'logo', 'report_title', 'group', 'topic',
      'instructor', 'student', 'class_info', 'year', 'empty', 'other'
    """
    text = para.text.strip()
    text_upper = text.upper()

    # 1. Đoạn trống có hình ảnh → logo
    if not text and has_visual_elements(para):
        return 'logo'

    # 2. Đoạn hoàn toàn trống → đánh dấu xóa
    if not text:
        return 'empty'

    # 3. Tên trường / bộ / khoa / viện
    if any(kw in text_upper for kw in HEADER_KW):
        return 'header'

    # 4. Tiêu đề loại báo cáo (BÁO CÁO MÔN HỌC, ĐỒ ÁN TỐT NGHIỆP...)
    if any(kw in text_upper for kw in TITLE_KW):
        return 'report_title'

    # 5. Nhóm / tổ
    if 'NHÓM' in text_upper or 'TỔ' in text_upper:
        return 'group'

    # 6. Giáo viên / giảng viên hướng dẫn
    if any(kw in text_upper for kw in ['GIÁO VIÊN', 'GIẢNG VIÊN', 'HƯỚNG DẪN', 'CÁN BỘ']):
        return 'instructor'

    # 7. Sinh viên / học viên thực hiện
    if any(kw in text_upper for kw in ['SINH VIÊN', 'HỌC VIÊN', 'THỰC HIỆN', 'NGƯỜI THỰC HIỆN']):
        return 'student'

    # 8. Lớp
    if 'LỚP' in text_upper and ':' in text:
        return 'class_info'

    # 9. Năm học / học kỳ
    if any(kw in text_upper for kw in ['NĂM HỌC', 'HỌC KỲ', 'NIÊN KHÓA', 'KHÓA']):
        return 'year'

    # 10. Chủ đề / đề tài (font lớn, thường >= 16pt)
    if max_run_font_size(para) >= 16:
        return 'topic'

    return 'other'


def apply_cover_spacing(para, cls, is_last_header=False):
    """Đặt spacing cho 1 phần tử trang bìa dựa trên phân loại."""
    # Bảng spacing: (before_pt, after_pt, line_spacing)
    SPACING = {
        'header':       (0,  2,  1.15),
        'logo':         (0,  18, 1.0),
        'report_title': (12, 6,  1.2),
        'group':        (6,  6,  1.15),
        'topic':        (6,  24, 1.2),
        'instructor':   (0,  6,  1.15),
        'student':      (0,  6,  1.15),
        'class_info':   (24, 0,  1.15),
        'year':         (6,  0,  1.15),
        'other':        (0,  0,  1.0),
    }

    before, after, ls = SPACING.get(cls, (0, 0, 1.0))

    # Header cuối cùng (thường là dòng KHOA...) cần khoảng cách lớn hơn
    if cls == 'header' and is_last_header:
        after = 12

    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if ls == 1.0:
        pf.line_spacing      = 1.0
        pf.line_spacing_rule  = WD_LINE_SPACING.SINGLE
    else:
        pf.line_spacing      = ls
        pf.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE

    # Bìa: không thụt đầu dòng, căn giữa (nếu chưa set)
    pf.first_line_indent = Cm(0)
    pf.left_indent       = Cm(0)
    if para.alignment is None:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ╔══════════════════════════════════════════════════════════════╗
# ║ TOC DETECTION                                               ║
# ╚══════════════════════════════════════════════════════════════╝

def has_existing_toc(doc):
    """
    Kiểm tra document đã có mục lục tự động chưa.
    Tìm: (a) đoạn có text 'MỤC LỤC' trước heading đầu tiên,
         (b) field code chứa 'TOC'.
    """
    # Cách 1: Tìm instrText chứa TOC
    body = doc.element.body
    for instrText in body.findall('.//' + qn('w:instrText')):
        if instrText.text and 'TOC' in instrText.text.upper():
            return True
    # Cách 2: Tìm đoạn 'MỤC LỤC' ở đầu document (trước heading đầu)
    for para in doc.paragraphs:
        if para.text.strip() == 'MỤC LỤC':
            return True
        if para.style and 'Heading' in para.style.name:
            break
    return False


def is_toc_paragraph(para):
    """Kiểm tra đoạn thuộc cấu trúc mục lục (field code hoặc heading MỤC LỤC)."""
    if para.text.strip() == 'MỤC LỤC':
        return True
    for instrText in para._element.findall('.//' + qn('w:instrText')):
        if instrText.text and 'TOC' in instrText.text.upper():
            return True
    # Kiểm tra fldChar (các run có field char thuộc TOC)
    for fldChar in para._element.findall('.//' + qn('w:fldChar')):
        return True
    return False


# ╔══════════════════════════════════════════════════════════════╗
# ║ MAIN FORMATTING FUNCTION                                    ║
# ╚══════════════════════════════════════════════════════════════╝

def format_document(input_path, output_path=None, log=print, settings=None):
    """
    Căn chỉnh toàn diện 1 file báo cáo đại học Việt Nam.
    
    Args:
        input_path:  Đường dẫn file .docx gốc
        output_path: Đường dẫn file kết quả (None → tự tạo tên)
        log:         Hàm in thông báo (mặc định print)
        settings:    Dict tùy chỉnh thông số (None → dùng mặc định)
    
    Returns:
        Đường dẫn file kết quả
    
    Raises:
        FileNotFoundError: File không tồn tại
        ValueError:        File không phải .docx
    """
    # ── User-configurable settings ──
    s = settings or {}
    font    = s.get('font_name',    FONT_NAME)
    bsz     = float(s.get('body_size',    BODY_SIZE_PT))
    h1sz    = float(s.get('h1_size',      H1_SIZE_PT))
    lspc    = float(s.get('line_spacing',  LINE_SPACING_MUL))
    fidt    = float(s.get('first_indent',  FIRST_INDENT_CM))
    m_l     = float(s.get('margin_left',   MARGIN_LEFT_MM))
    m_r     = float(s.get('margin_right',  MARGIN_RIGHT_MM))
    m_t     = float(s.get('margin_top',    MARGIN_TOP_MM))
    m_b     = float(s.get('margin_bot',    MARGIN_BOT_MM))

    # ── Validate input ──
    input_path = str(input_path)
    if not os.path.isfile(input_path):
        raise FileNotFoundError(f"Không tìm thấy file: {input_path}")
    if not input_path.lower().endswith('.docx'):
        raise ValueError(f"File phải có đuôi .docx: {input_path}")

    # ── Generate output path ──
    if output_path is None:
        p = Path(input_path)
        base = p.stem
        for suffix in [' - ĐÃ CHỈNH SỬA', ' - DA CHINH SUA']:
            while base.endswith(suffix):
                base = base[:-len(suffix)]
        out_name = f"{base} - ĐÃ CHỈNH SỬA{p.suffix}"
        output_path = str(p.parent / out_name)

    doc = Document(input_path)
    total_paras = len(doc.paragraphs)
    total_tables = len(doc.tables)
    log(f"  Đã mở file: {os.path.basename(input_path)}")
    log(f"  Tổng: {total_paras} đoạn văn, {total_tables} bảng")

    # ════════════════════════════════════════════════════════════
    # BƯỚC 1: THIẾT LẬP LỀ TRANG (A4)
    # ════════════════════════════════════════════════════════════
    log("  [1/7] Thiết lập lề trang A4...")
    # Tính chiều rộng nội dung cho TOC tab stop
    toc_right_tab_cm = (210 - m_l - m_r) / 10.0
    for section in doc.sections:
        section.page_width   = Mm(210)
        section.page_height  = Mm(297)
        section.left_margin  = Mm(m_l)
        section.right_margin = Mm(m_r)
        section.top_margin   = Mm(m_t)
        section.bottom_margin= Mm(m_b)

    # ════════════════════════════════════════════════════════════
    # BƯỚC 2: SỐ TRANG (ẩn trang bìa)
    # ════════════════════════════════════════════════════════════
    log("  [2/7] Thiết lập số trang...")
    for section in doc.sections:
        section.different_first_page_header_footer = True

        # Footer thường (trang 2+): số trang căn giữa
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after  = Pt(0)

        # Chèn field PAGE
        fp.add_run()._element.append(
            parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        fp.add_run()._element.append(
            parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        fp.add_run()._element.append(
            parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
        fp.add_run("1")
        fp.add_run()._element.append(
            parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
        for r in fp.runs:
            r.font.name = font
            r.font.size = Pt(bsz)

        # Footer trang đầu: để trống
        fpf = section.first_page_footer
        fpf.is_linked_to_previous = False
        for p in fpf.paragraphs:
            p.clear()

    # ════════════════════════════════════════════════════════════
    # BƯỚC 3: ĐỊNH DẠNG STYLES
    # ════════════════════════════════════════════════════════════
    log("  [3/7] Định dạng kiểu chữ (styles)...")

    # ── Normal ──
    ns = doc.styles['Normal']
    ns.font.name = font
    ns.font.size = Pt(bsz)
    ns.paragraph_format.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
    ns.paragraph_format.first_line_indent   = Cm(fidt)
    ns.paragraph_format.left_indent         = Cm(0)
    ns.paragraph_format.line_spacing        = lspc
    ns.paragraph_format.line_spacing_rule   = WD_LINE_SPACING.MULTIPLE
    ns.paragraph_format.space_before        = Pt(3)
    ns.paragraph_format.space_after         = Pt(3)
    ns.paragraph_format.page_break_before   = False
    set_fonts_xml(ns, font)

    # ── Heading 1 / 2 / 3 ──
    for sname, sz, align, caps in [
        ('Heading 1', h1sz,  WD_ALIGN_PARAGRAPH.CENTER,  True),
        ('Heading 2', bsz,   WD_ALIGN_PARAGRAPH.JUSTIFY, False),
        ('Heading 3', bsz,   WD_ALIGN_PARAGRAPH.JUSTIFY, False),
    ]:
        try:
            s = doc.styles[sname]
        except KeyError:
            continue
        s.font.name     = font
        s.font.size     = Pt(sz)
        s.font.bold     = True
        s.font.all_caps = caps
        s.paragraph_format.alignment        = align
        s.paragraph_format.first_line_indent = Cm(0)
        s.paragraph_format.left_indent       = Cm(0)
        s.paragraph_format.line_spacing      = lspc
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        s.paragraph_format.keep_with_next    = True
        s.paragraph_format.page_break_before = (sname == 'Heading 1')
        if sname == 'Heading 1':
            s.paragraph_format.space_before = Pt(12)
            s.paragraph_format.space_after  = Pt(6)
        else:
            s.paragraph_format.space_before = Pt(6)
            s.paragraph_format.space_after  = Pt(3)
        set_fonts_xml(s, font)

    # ── List Paragraph ──
    try:
        lps = doc.styles['List Paragraph']
        lps.font.name = font
        lps.font.size = Pt(bsz)
        lps.paragraph_format.alignment        = WD_ALIGN_PARAGRAPH.JUSTIFY
        lps.paragraph_format.line_spacing      = lspc
        lps.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        lps.paragraph_format.space_before      = Pt(3)
        lps.paragraph_format.space_after       = Pt(3)
        lps.paragraph_format.page_break_before = False
        set_fonts_xml(lps, font)
    except KeyError:
        pass

    # ── TOC 1 / 2 / 3 ──
    for lv in range(1, 4):
        tn = f'TOC {lv}'
        try:
            ts = doc.styles[tn]
        except KeyError:
            ts = doc.styles.add_style(tn, WD_STYLE_TYPE.PARAGRAPH)

        if ts._element.get(qn('w:customStyle')) is not None:
            del ts._element.attrib[qn('w:customStyle')]

        ts.font.name = font
        ts.font.size = Pt(bsz)
        ts.font.bold = (lv == 1)
        pf = ts.paragraph_format
        pf.line_spacing      = lspc
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.space_before      = Pt(3)
        pf.space_after       = Pt(3)
        pf.page_break_before = False
        pf.right_indent      = Cm(1.0)

        if lv == 1:
            pf.left_indent       = Cm(0)
            pf.first_line_indent = Cm(0)
        elif lv == 2:
            pf.left_indent       = Cm(1.5)
            pf.first_line_indent = Cm(-0.7)
        else:
            pf.left_indent       = Cm(2.2)
            pf.first_line_indent = Cm(-0.7)

        pPr = ts._element.find(qn('w:pPr'))
        if pPr is not None:
            old_tabs = pPr.find(qn('w:tabs'))
            if old_tabs is not None:
                pPr.remove(old_tabs)
        pf.tab_stops.add_tab_stop(
            Cm(toc_right_tab_cm),
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.DOTS,
        )
        set_fonts_xml(ts, font)

    # ════════════════════════════════════════════════════════════
    # BƯỚC 4: ĐỊNH DẠNG TỪNG ĐOẠN VĂN
    # ════════════════════════════════════════════════════════════
    log("  [4/7] Định dạng nội dung...")

    cover_end = find_cover_end(doc)
    log(f"         Trang bìa: {cover_end} đoạn")

    in_references = False
    paragraphs_to_delete = []

    # ── Phân loại các phần tử trang bìa ──
    cover_classes = []
    for i in range(min(cover_end, len(doc.paragraphs))):
        cover_classes.append(classify_cover_element(doc.paragraphs[i]))

    # Xác định header cuối cùng (cần space lớn hơn)
    last_header_idx = -1
    for ci, cls in enumerate(cover_classes):
        if cls == 'header':
            last_header_idx = ci

    # ── Duyệt từng đoạn ──
    for i, para in enumerate(doc.paragraphs):
        text       = para.text.strip()
        style_name = para.style.name if para.style else 'Normal'
        pPr        = para._element.find(qn('w:pPr'))

        # ────────────────────────────────────────
        # A. TRANG BÌA
        # ────────────────────────────────────────
        if i < cover_end:
            cls = cover_classes[i] if i < len(cover_classes) else 'other'

            # Xóa đoạn trống (giữ đoạn có hình)
            if cls == 'empty':
                paragraphs_to_delete.append(para)
                continue

            # Đặt spacing theo phân loại
            is_last_hdr = (cls == 'header' and i == last_header_idx)
            apply_cover_spacing(para, cls, is_last_header=is_last_hdr)

            # Font
            for run in para.runs:
                set_run_fonts(run, font)
            continue

        # ────────────────────────────────────────
        # B. BỎ QUA ĐOẠN MỤC LỤC ĐÃ CÓ
        # ────────────────────────────────────────
        if is_toc_paragraph(para):
            continue

        # ────────────────────────────────────────
        # C. XÓA ĐOẠN TRỐNG TRONG NỘI DUNG
        # ────────────────────────────────────────
        is_empty = not text
        has_img  = has_visual_elements(para)
        if is_empty and not has_img:
            paragraphs_to_delete.append(para)
            continue

        # ────────────────────────────────────────
        # D. XÓA NGẮT TRANG THỦ CÔNG TRONG RUN
        # ────────────────────────────────────────
        for run in para.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type'), '') == 'page':
                    run._element.remove(br)

        # ────────────────────────────────────────
        # E. THEO DÕI VÙNG TÀI LIỆU THAM KHẢO
        # ────────────────────────────────────────
        if text in REF_TITLES:
            in_references = True
        elif style_name == 'Heading 1' and text not in REF_TITLES:
            in_references = False

        # ────────────────────────────────────────
        # F. TIÊU ĐỀ ĐẶC BIỆT → Heading 1 không số
        # ────────────────────────────────────────
        if text in SPECIAL_TITLES:
            para.style = doc.styles['Heading 1']
            disable_numbering(para)
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.left_indent       = Cm(0)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_fonts(run, font)
                run.font.size = Pt(h1sz)
            continue

        # ────────────────────────────────────────
        # G. HEADING (giữ auto-numbering)
        # ────────────────────────────────────────
        if 'Heading' in style_name:
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.left_indent       = Cm(0)
            if style_name == 'Heading 1':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs:
                    r.font.size = Pt(h1sz)
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in para.runs:
                    r.font.size = Pt(bsz)
            for r in para.runs:
                set_run_fonts(r, font)

        # ────────────────────────────────────────
        # H. LIST PARAGRAPH / ĐOẠN CÓ NUMBERING
        # ────────────────────────────────────────
        elif (style_name == 'List Paragraph' or
              (pPr is not None and pPr.find(qn('w:numPr')) is not None)):
            has_num = False
            ilvl = 0
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    has_num = True
                    ilvl_el = numPr.find(qn('w:ilvl'))
                    if ilvl_el is not None:
                        try:
                            ilvl = int(ilvl_el.get(qn('w:val')))
                        except (ValueError, TypeError):
                            ilvl = 0
            if has_num:
                para.paragraph_format.left_indent       = Cm(1.27 + ilvl * 0.63)
                para.paragraph_format.first_line_indent = Cm(-0.63)
            else:
                para.paragraph_format.first_line_indent = Cm(fidt)
                para.paragraph_format.left_indent       = Cm(0)

            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in para.runs:
                set_run_fonts(r, font)
                r.font.size = Pt(bsz)

        # ────────────────────────────────────────
        # I. ĐOẠN BÌNH THƯỜNG (Normal)
        # ────────────────────────────────────────
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Cm(fidt)
            para.paragraph_format.left_indent       = Cm(0)
            for r in para.runs:
                set_run_fonts(r, font)
                r.font.size = Pt(bsz)

        # ── Tài liệu tham khảo: thụt treo ──
        if in_references and 'Heading' not in style_name:
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.left_indent       = Cm(1)

        # ── Spacing cho nội dung (không áp dụng lên heading) ──
        if 'Heading' not in style_name:
            para.paragraph_format.line_spacing      = lspc
            para.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.space_before       = Pt(3)
            para.paragraph_format.space_after        = Pt(3)

    # ── Xóa đoạn trống ──
    log(f"         Xóa {len(paragraphs_to_delete)} đoạn trống...")
    for para in paragraphs_to_delete:
        el = para._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    # ════════════════════════════════════════════════════════════
    # BƯỚC 5: ĐỊNH DẠNG BẢNG
    # ════════════════════════════════════════════════════════════
    log(f"  [5/7] Định dạng {total_tables} bảng...")

    # Xác định bảng nào nằm trên trang bìa
    # Logic: bảng xuất hiện trong body XML trước đoạn cover_end → bảng bìa
    cover_table_indices = set()
    if cover_end > 0:
        body_el = doc.element.body
        cover_para_elements = set()
        para_count = 0
        for child in body_el:
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                if para_count >= cover_end:
                    break
                cover_para_elements.add(id(child))
                para_count += 1
            elif tag == 'tbl':
                if para_count < cover_end:
                    # Bảng này nằm trong vùng trang bìa
                    for ti, t in enumerate(doc.tables):
                        if id(t._tbl) == id(child):
                            cover_table_indices.add(ti)
                            break

    for ti, table in enumerate(doc.tables):
        is_cover_table = ti in cover_table_indices

        if is_cover_table:
            # ── Bảng trang bìa (danh sách SV): không viền, giữ numbering ──
            for row in table.rows:
                trPr = row._tr.get_or_add_trPr()
                # Không tách dòng qua trang
                if trPr.find(qn('w:cantSplit')) is None:
                    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para.paragraph_format.line_spacing     = 1.15
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        para.paragraph_format.space_before      = Pt(0)
                        para.paragraph_format.space_after       = Pt(0)
                        for r in para.runs:
                            set_run_fonts(r, font)
            continue

        # ── Bảng nội dung: viền đầy đủ, header đậm căn giữa ──
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)

        # Rộng 100%
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblPr.append(parse_xml(
                f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))
        else:
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')

        # Viền
        old_borders = tblPr.find(qn('w:tblBorders'))
        if old_borders is not None:
            tblPr.remove(old_borders)
        tblPr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'</w:tblBorders>'
        ))

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    pf = para.paragraph_format
                    pf.first_line_indent  = Cm(0)
                    pf.left_indent        = Cm(0)
                    pf.space_before       = Pt(2)
                    pf.space_after        = Pt(2)
                    pf.line_spacing       = lspc
                    pf.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
                    pf.page_break_before  = False
                    para.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                                      if row_idx == 0
                                      else WD_ALIGN_PARAGRAPH.JUSTIFY)
                    for r in para.runs:
                        set_run_fonts(r, font)
                        r.font.size = Pt(bsz)
                        if row_idx == 0:
                            r.font.bold = True

                # Căn giữa dọc
                tc   = cell._element
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                    tc.insert(0, tcPr)
                vA = tcPr.find(qn('w:vAlign'))
                if vA is None:
                    tcPr.append(parse_xml(
                        f'<w:vAlign {nsdecls("w")} w:val="center"/>'))
                else:
                    vA.set(qn('w:val'), 'center')

    # ════════════════════════════════════════════════════════════
    # BƯỚC 6: CHÈN MỤC LỤC TỰ ĐỘNG
    # ════════════════════════════════════════════════════════════
    log("  [6/7] Tạo mục lục...")

    if has_existing_toc(doc):
        log("         ⏩ Mục lục đã tồn tại — bỏ qua.")
    else:
        # Tìm vị trí chèn: trước Heading đầu tiên hoặc tiêu đề đặc biệt đầu tiên
        insert_before = None
        for para in doc.paragraphs:
            style = para.style.name if para.style else ''
            text  = para.text.strip()
            if 'Heading' in style:
                insert_before = para._element
                break
            if text in SPECIAL_TITLES:
                insert_before = para._element
                break

        if insert_before is not None:
            body = doc.element.body
            idx  = list(body).index(insert_before)

            # Heading MỤC LỤC
            body.insert(idx, parse_xml(
                f'<w:p {nsdecls("w")}>'
                f'<w:pPr>'
                f'<w:pageBreakBefore/>'
                f'<w:jc w:val="center"/>'
                f'<w:spacing w:before="240" w:after="120"/>'
                f'</w:pPr>'
                f'<w:r>'
                f'<w:rPr>'
                f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" '
                f'w:eastAsia="{font}" w:cs="{font}"/>'
                f'<w:b/>'
                f'<w:sz w:val="32"/><w:szCs w:val="32"/>'
                f'</w:rPr>'
                f'<w:t>MỤC LỤC</w:t>'
                f'</w:r>'
                f'</w:p>'
            ))
            idx += 1

            # TOC field
            body.insert(idx, parse_xml(
                f'<w:p {nsdecls("w")}>'
                f'<w:pPr><w:spacing w:before="60" w:after="60"/></w:pPr>'
                f'<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
                f'<w:r><w:instrText xml:space="preserve">'
                f' TOC \\o "1-3" \\h \\z \\u '
                f'</w:instrText></w:r>'
                f'<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
                f'<w:r><w:rPr>'
                f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}"/>'
                f'<w:sz w:val="26"/>'
                f'</w:rPr>'
                f'<w:t>Nhấn Ctrl+A rồi F9 để cập nhật mục lục</w:t></w:r>'
                f'<w:r><w:fldChar w:fldCharType="end"/></w:r>'
                f'</w:p>'
            ))
            log("         ✅ Đã chèn mục lục tự động.")
        else:
            log("         ⚠️ Không tìm thấy vị trí chèn mục lục.")

    # Yêu cầu Word cập nhật mục lục khi mở file
    try:
        settings = doc.settings.element
        uf = settings.find(qn('w:updateFields'))
        if uf is None:
            settings.append(parse_xml(
                f'<w:updateFields {nsdecls("w")} w:val="true"/>'))
        else:
            uf.set(qn('w:val'), 'true')
    except Exception:
        pass

    # ════════════════════════════════════════════════════════════
    # BƯỚC 7: LƯU FILE
    # ════════════════════════════════════════════════════════════
    log("  [7/7] Lưu file...")
    doc.save(output_path)
    log(f"  ✅ Hoàn thành: {output_path}")

    return output_path


# ╔══════════════════════════════════════════════════════════════╗
# ║ CLI ENTRY POINT                                             ║
# ╚══════════════════════════════════════════════════════════════╝

def main():
    """Xử lý khi chạy từ command line / kéo thả vào .bat"""
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')

    banner = """
  ╔══════════════════════════════════════════════════════════╗
  ║       BOT CĂN CHỈNH BÁO CÁO TỰ ĐỘNG  v1.0           ║
  ╠══════════════════════════════════════════════════════════╣
  ║  Lề: Trái 3cm │ Phải 1.5cm │ Trên/Dưới 2cm           ║
  ║  Font: Times New Roman  │  Giãn dòng: 1.3             ║
  ║  Heading 1: 16pt đậm căn giữa IN HOA                  ║
  ║  Heading 2/3: 13pt đậm căn đều                        ║
  ║  Nội dung: 13pt căn đều, thụt 1cm                     ║
  ╚══════════════════════════════════════════════════════════╝
    """

    if len(sys.argv) < 2:
        print(banner)
        print("  Cách sử dụng:")
        print("    Kéo thả file .docx vào file .bat trên Desktop")
        print("    Hoặc: python format_engine.py <file.docx>")
        print()
        input("  Nhấn Enter để đóng...")
        return

    print(banner)

    success_count = 0
    fail_count = 0

    for filepath in sys.argv[1:]:
        filepath = filepath.strip('"')
        print(f"\n  ▶ Xử lý: {os.path.basename(filepath)}")
        print("  " + "─" * 50)
        try:
            output = format_document(filepath)
            success_count += 1
        except FileNotFoundError as e:
            print(f"  ❌ Lỗi: {e}")
            fail_count += 1
        except ValueError as e:
            print(f"  ❌ Lỗi: {e}")
            fail_count += 1
        except Exception as e:
            print(f"  ❌ Lỗi không xác định: {e}")
            traceback.print_exc()
            fail_count += 1

    print("\n  " + "═" * 50)
    print(f"  Kết quả: {success_count} thành công, {fail_count} thất bại")
    print()


if __name__ == '__main__':
    main()
